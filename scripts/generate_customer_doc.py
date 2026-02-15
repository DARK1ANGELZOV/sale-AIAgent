from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


def set_base_font(document: Document, font_name: str = "Times New Roman", size: int = 12) -> None:
    style = document.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(size)
    style._element.rPr.rFonts.set(qn("w:ascii"), font_name)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    style._element.rPr.rFonts.set(qn("w:cs"), font_name)


def add_heading(document: Document, text: str, level: int = 1) -> None:
    paragraph = document.add_heading(text, level=level)
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run._element.rPr.rFonts.set(qn("w:cs"), "Times New Roman")


def add_paragraph(document: Document, text: str, bold: bool = False) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:cs"), "Times New Roman")
    paragraph.paragraph_format.space_after = Pt(6)


def add_list(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Number")
        run = paragraph.add_run(item)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run._element.rPr.rFonts.set(qn("w:cs"), "Times New Roman")
        paragraph.paragraph_format.space_after = Pt(4)


def add_code_block(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run._element.rPr.rFonts.set(qn("w:cs"), "Consolas")
    paragraph.paragraph_format.space_after = Pt(8)


def build_document() -> Document:
    doc = Document()
    set_base_font(doc)

    section = doc.sections[0]
    section.top_margin = Pt(56)
    section.bottom_margin = Pt(56)
    section.left_margin = Pt(64)
    section.right_margin = Pt(64)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("ТЕХНИЧЕСКИЙ ДОКУМЕНТ")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:cs"), "Times New Roman")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run("Пилотный прототип ИИ-агента для поддержки продаж и технических подразделений")
    sub_run.font.size = Pt(13)
    sub_run.font.name = "Times New Roman"
    sub_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    sub_run._element.rPr.rFonts.set(qn("w:cs"), "Times New Roman")

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    m_run = meta.add_run(f"Дата: {date.today().strftime('%d.%m.%Y')}   |   Версия: 1.0")
    m_run.font.name = "Times New Roman"
    m_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    m_run._element.rPr.rFonts.set(qn("w:cs"), "Times New Roman")

    doc.add_page_break()

    add_heading(doc, "1. Общие положения", level=1)
    add_paragraph(
        doc,
        "Настоящий документ определяет технические требования к пилотному прототипу ИИ-агента, "
        "предназначенному для предоставления проверенной информации по одному продукту компании "
        "на основе ограниченного набора документов.",
    )
    add_paragraph(doc, "Цель пилота — подтвердить техническую реализуемость решения и качество ответов.")
    add_list(
        doc,
        [
            "Подтвердить корректность архитектуры RAG (Retrieval-Augmented Generation).",
            "Подтвердить корректность ответов при обязательном цитировании источников.",
            "Подтвердить устойчивую работу без обучения собственной LLM и без использования GPU.",
        ],
    )
    add_paragraph(doc, "Ограничения проекта:", bold=True)
    add_list(
        doc,
        [
            "Обучение собственной LLM не допускается.",
            "Использование специализированных GPU не допускается.",
            "Применяются только готовые модели через API.",
            "Ответ формируется исключительно на основе загруженных документов.",
            "Ответ без ссылок на источник считается некорректным.",
        ],
    )

    add_heading(doc, "2. Входные данные", level=1)
    add_list(
        doc,
        [
            "Набор из 10–20 документов.",
            "Поддерживаемые форматы: PDF, Excel, Word.",
            "Материалы по одному продукту или продуктовой линейке.",
            "Тестовые вопросы двух типов: продажи и технические.",
        ],
    )

    add_heading(doc, "3. Архитектура решения (RAG)", level=1)
    add_paragraph(doc, "Архитектура включает ingestion, retrieval, генерацию ответа и контроль корректности.")
    add_paragraph(doc, "Текстовая блок-схема:", bold=True)
    add_code_block(
        doc,
        "[Web/Telegram UI]\n"
        "        |\n"
        "        v\n"
        "[FastAPI API Layer]\n"
        "        |\n"
        "        v\n"
        "[Request Orchestrator] ---------------------------+\n"
        "        |                                         |\n"
        "        |                                         v\n"
        "        |                                 [JSON Logging]\n"
        "        v\n"
        "[Retriever]\n"
        "  - Query preprocessing\n"
        "  - Query embedding\n"
        "  - Qdrant top-k retrieval\n"
        "  - Re-ranking + threshold\n"
        "        |\n"
        "        v\n"
        "[Context Builder]\n"
        "        |\n"
        "        v\n"
        "[LLM API]\n"
        "        |\n"
        "        v\n"
        "[Post-validation]\n"
        "  - Проверка наличия цитат\n"
        "  - Проверка соответствия цитаты источнику\n"
        "  - Политика отказа при низкой релевантности\n"
        "        |\n"
        "        v\n"
        "[Final Response + Sources]"
    )

    add_paragraph(doc, "Пайплайн загрузки и индексации:", bold=True)
    add_code_block(
        doc,
        "[Upload] -> [PDF/Excel/Word parsers] -> [Table normalization] -> "
        "[Chunking + metadata] -> [Embeddings] -> [Qdrant Index]"
    )

    add_heading(doc, "4. Выбор LLM (через API)", level=1)
    add_paragraph(doc, "Выбранная модель: gpt-4o-mini.")
    add_paragraph(doc, "Обоснование:", bold=True)
    add_list(
        doc,
        [
            "Достаточное качество для RAG-сценариев на русском языке.",
            "Стабильная скорость ответа для интерактивного чата.",
            "Экономичная стоимость на пилотном объеме запросов.",
            "Поддержка строгого формата ответа с обязательными источниками.",
        ],
    )

    add_heading(doc, "5. Embeddings и векторная база данных", level=1)
    add_paragraph(doc, "Модель embeddings: sentence-transformers all-MiniLM-L6-v2, размерность 384.")
    add_list(
        doc,
        [
            "CPU-совместимость и низкие требования к ресурсам.",
            "Достаточная точность семантического поиска для пилота 10–20 документов.",
            "Пакетная обработка embeddings для ускорения индексации.",
        ],
    )
    add_paragraph(doc, "Векторная база: Qdrant.")
    add_list(
        doc,
        [
            "Поддержка поиска top-k и фильтрации по metadata (включая version).",
            "Подходит для пилота и дальнейшего масштабирования.",
            "Простая эксплуатация и предсказуемая производительность на CPU.",
        ],
    )

    add_heading(doc, "6. Chunking, таблицы и метаданные", level=1)
    add_paragraph(doc, "Применяется логически ориентированная сегментация текста.")
    add_list(
        doc,
        [
            "Размер чанка: 800–1200 символов.",
            "Overlap: 120–200 символов.",
            "Границы чанков учитывают разделы и заголовки документа.",
            "Таблицы выделяются в отдельные смысловые блоки.",
        ],
    )
    add_paragraph(doc, "Метаданные чанка:", bold=True)
    add_list(
        doc,
        [
            "document_name",
            "document_id",
            "version",
            "page_number (для PDF)",
            "section (для Word/PDF)",
            "sheet_name (для Excel)",
            "chunk_id",
            "timestamp",
        ],
    )
    add_paragraph(doc, "Работа с Excel и таблицами PDF:", bold=True)
    add_list(
        doc,
        [
            "Excel обрабатывается по листам и строкам с сохранением контекста заголовков.",
            "Таблицы из PDF нормализуются в текстовые строки вида «параметр = значение».",
            "Числовые и сравнительные запросы выполняются по извлеченным табличным данным.",
            "Для каждого числового вывода обязательно указывается источник.",
        ],
    )

    add_heading(doc, "7. Версионность и обновление данных", level=1)
    add_paragraph(doc, "В системе реализован версионный учет документов.")
    add_list(
        doc,
        [
            "Каждому документу присваивается постоянный document_id.",
            "Каждая новая загрузка формирует новую version.",
            "Предыдущие версии сохраняются как soft delete (is_active=false).",
            "Поиск по умолчанию выполняется по активной версии.",
        ],
    )
    add_paragraph(doc, "Пайплайн обновления:", bold=True)
    add_list(
        doc,
        [
            "Проверка контрольной суммы файла.",
            "Повторный парсинг, chunking и расчет embeddings при изменении.",
            "Переиндексация в Qdrant и атомарное переключение активной версии.",
            "Проверка целостности индекса и доступности retrieval.",
        ],
    )

    add_heading(doc, "8. Ограничение галлюцинаций и формат ответа", level=1)
    add_list(
        doc,
        [
            "В LLM передается только извлеченный контекст.",
            "При релевантности ниже порога возвращается отказ.",
            "Ответ без источников блокируется пост-валидацией.",
            "Цитата проверяется на дословное соответствие retrieved-контексту.",
        ],
    )
    add_paragraph(doc, "Шаблон ответа:", bold=True)
    add_code_block(
        doc,
        "Ответ:\n"
        "<текст ответа>\n\n"
        "Источники:\n"
        "1. Документ: <имя документа>\n"
        "   Страница/раздел: <значение>\n"
        "   Цитата: \"<точный фрагмент>\""
    )
    add_paragraph(doc, "Сообщение отказа при отсутствии данных:", bold=True)
    add_code_block(doc, "В базе знаний отсутствуют релевантные данные.")

    add_heading(doc, "9. Инфраструктурные требования (без GPU)", level=1)
    add_paragraph(doc, "Минимальная конфигурация (API-версия):")
    add_list(
        doc,
        [
            "CPU: 4 vCPU",
            "RAM: 8 GB",
            "Диск: 40 GB SSD",
            "GPU: не требуется",
        ],
    )
    add_paragraph(doc, "Рекомендуемая конфигурация:")
    add_list(
        doc,
        [
            "CPU: 8 vCPU",
            "RAM: 16 GB",
            "Диск: 80 GB SSD",
            "GPU: не требуется",
        ],
    )

    add_heading(doc, "10. Оценка стоимости эксплуатации", level=1)
    add_paragraph(
        doc,
        "Расчет выполнен для API-подключения при среднем запросе: "
        "2500 входных токенов и 550 выходных токенов.",
    )
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Объем"
    hdr[1].text = "Входные токены"
    hdr[2].text = "Выходные токены"
    hdr[3].text = "Стоимость input"
    hdr[4].text = "Стоимость output"
    hdr[5].text = "Итого"

    rows = [
        ("100 запросов", "250 000", "55 000", "$0.0375", "$0.0330", "$0.0705"),
        ("1000 запросов", "2 500 000", "550 000", "$0.3750", "$0.3300", "$0.7050"),
    ]
    for values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            cells[idx].text = value

    add_paragraph(
        doc,
        "Примечание: фактическая стоимость зависит от реальной длины контекста и актуальных тарифов API-провайдера.",
    )

    add_heading(doc, "11. Логирование", level=1)
    add_paragraph(doc, "Формат хранения: JSON Lines (одна запись на запрос).")
    add_list(
        doc,
        [
            "Текст запроса",
            "Дата/время",
            "Тип запроса (sales/technical)",
            "Использованные документы и версии",
            "Confidence score",
            "Время ответа",
            "Ошибки и статус обработки",
        ],
    )
    add_paragraph(doc, "Пример JSON-записи:")
    add_code_block(
        doc,
        "{\n"
        '  "timestamp": "2026-02-15T13:45:21Z",\n'
        '  "request_id": "a7c9c2d1-xxxx",\n'
        '  "query_type": "technical",\n'
        '  "question": "Какие ключевые функции продукта?",\n'
        '  "used_documents": [{"document_name": "АстроСекьюр_5000_Техническая_спецификация.docx", "version": "1.2"}],\n'
        '  "confidence": 0.87,\n'
        '  "processing_time_ms": 1480,\n'
        '  "status": "ok"\n'
        "}"
    )

    add_heading(doc, "12. Демонстрация и критерии оценки", level=1)
    add_paragraph(doc, "Исполнитель демонстрирует:")
    add_list(
        doc,
        [
            "Загрузку и индексацию документов PDF/Excel/Word.",
            "Ответы на продажи и технические вопросы с обязательными цитатами.",
            "Корректный отказ при отсутствии релевантных данных.",
            "Работу с числовыми и сравнительными вопросами по таблицам.",
            "Масштабируемость и план дальнейшего развития решения.",
        ],
    )
    add_paragraph(doc, "Критерии оценки пилота:")
    add_list(
        doc,
        [
            "Точность и полнота ответов.",
            "Корректность цитирования.",
            "Отсутствие ответов вне базы знаний.",
            "Обоснованность архитектурных решений.",
            "Экономическая эффективность эксплуатации.",
        ],
    )

    add_paragraph(doc, "Документ подготовлен для предоставления заказчику.", bold=True)
    return doc


def main() -> None:
    output_dir = Path("deliverables")
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / "Технический_документ_Пилот_RAG_для_заказчика.docx"

    document = build_document()
    document.save(output_path)
    print(output_path.resolve())


if __name__ == "__main__":
    main()
